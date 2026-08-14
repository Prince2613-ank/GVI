import re, os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

SRC = r"c:\Users\princ\Downloads\GVI\GVI_PRODUCT_DOCUMENTATION.md"
OUT = r"c:\Users\princ\Downloads\GVI\GVI_Product_Documentation.docx"
IMG_DIR = r"c:\Users\princ\Downloads\GVI\docs\images"

GREEN = RGBColor(0x1B, 0x7F, 0x3B)
DARK = RGBColor(0x1A, 0x1A, 0x1A)
GREY = RGBColor(0x55, 0x55, 0x55)

missing_figures = []

doc = Document()

st = doc.styles["Normal"]
st.font.name = "Calibri"
st.font.size = Pt(11)
st.paragraph_format.space_after = Pt(8)
st.paragraph_format.line_spacing = 1.15

for name, size, color, bold in (
    ("Heading 1", 22, GREEN, True),
    ("Heading 2", 16, GREEN, True),
    ("Heading 3", 13, DARK, True),
    ("Heading 4", 11.5, DARK, True),
):
    s = doc.styles[name]
    s.font.name = "Calibri"
    s.font.size = Pt(size)
    s.font.color.rgb = color
    s.font.bold = bold
    s.paragraph_format.space_before = Pt(14)
    s.paragraph_format.space_after = Pt(6)


INLINE = re.compile(r"(\*\*.+?\*\*|\*[^*]+?\*|`[^`]+?`)")


def add_runs(par, text, bold=False, italic=False):
    """Render **bold**, *italic* and `code` as real Word runs.

    Recurses into bold spans so *italic* nested inside **bold** still renders
    rather than leaking literal asterisks into the document.
    """
    for token in INLINE.split(text):
        if not token:
            continue
        if token.startswith("**") and token.endswith("**") and len(token) > 4:
            add_runs(par, token[2:-2], bold=True, italic=italic)
        elif token.startswith("`") and token.endswith("`"):
            r = par.add_run(token[1:-1])
            r.font.name = "Consolas"
            r.font.size = Pt(10)
            if bold: r.bold = True
            if italic: r.italic = True
        elif token.startswith("*") and token.endswith("*") and len(token) > 2:
            add_runs(par, token[1:-1], bold=bold, italic=True)
        else:
            r = par.add_run(token)
            # Assigned only when true so heading/table styles keep their own
            # bold instead of being explicitly overridden to False.
            if bold: r.bold = True
            if italic: r.italic = True


def split_row(line):
    return [c.strip() for c in line.strip().strip("|").split("|")]


lines = open(SRC, encoding="utf-8").read().split("\n")
i = 0
while i < len(lines):
    line = lines[i].rstrip()
    stripped = line.strip()

    # --- figure: ![figure:file.png](caption) ---
    m = re.match(r"^!\[figure:([^\]]+)\]\((.*)\)\s*$", stripped)
    if m:
        filename, caption = m.group(1), m.group(2)
        path = os.path.join(IMG_DIR, filename)
        # A 0-byte file counts as missing — one of the supplied screenshots
        # arrived empty, and add_picture would raise rather than skip it.
        if os.path.exists(path) and os.path.getsize(path) > 0:
            pic = doc.add_paragraph()
            pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pic.paragraph_format.space_before = Pt(10)
            pic.paragraph_format.space_after = Pt(4)
            pic.add_run().add_picture(path, width=Inches(6.3))
        else:
            missing_figures.append(filename)
            ph = doc.add_paragraph()
            ph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = ph.add_run("[ image placeholder: %s ]" % filename)
            r.italic = True
            r.font.color.rgb = GREY
        if caption:
            cap = doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap.paragraph_format.space_after = Pt(14)
            add_runs(cap, caption)
            for r in cap.runs:
                r.italic = True
                r.font.size = Pt(9.5)
                r.font.color.rgb = GREY
        i += 1
        continue

    # --- table ---
    if stripped.startswith("|") and i + 1 < len(lines) and re.match(r"^\|[\s:|-]+\|$", lines[i + 1].strip()):
        header = split_row(stripped)
        i += 2
        rows = []
        while i < len(lines) and lines[i].strip().startswith("|"):
            rows.append(split_row(lines[i].strip()))
            i += 1
        table = doc.add_table(rows=1, cols=len(header))
        table.style = "Light Grid Accent 3"
        for c, text in enumerate(header):
            cell = table.rows[0].cells[c]
            cell.text = ""
            add_runs(cell.paragraphs[0], text)
            for r in cell.paragraphs[0].runs:
                r.bold = True
        for row in rows:
            cells = table.add_row().cells
            for c, text in enumerate(row[: len(header)]):
                cells[c].text = ""
                add_runs(cells[c].paragraphs[0], text)
        doc.add_paragraph()
        continue

    if not stripped or stripped == "---":
        i += 1
        continue

    # --- headings ---
    m = re.match(r"^(#{1,4})\s+(.*)", stripped)
    if m:
        level, text = len(m.group(1)), m.group(2)
        p = doc.add_paragraph(style=f"Heading {level}")
        add_runs(p, text)
        if level == 1:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        i += 1
        continue

    # --- blockquote ---
    if stripped.startswith(">"):
        body = stripped.lstrip("> ").strip()
        while i + 1 < len(lines) and lines[i + 1].strip().startswith(">"):
            i += 1
            body += " " + lines[i].strip().lstrip("> ").strip()
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.4)
        add_runs(p, body)
        for r in p.runs:
            r.italic = True
            r.font.color.rgb = GREEN
        i += 1
        continue

    # --- list items (with continuation lines) ---
    m = re.match(r"^(\s*)([-*]|\d+\.)\s+(.*)", line)
    if m:
        indent, marker, text = m.group(1), m.group(2), m.group(3)
        while i + 1 < len(lines):
            nxt = lines[i + 1]
            if nxt.strip() and not re.match(r"^\s*([-*]|\d+\.)\s+", nxt) and not nxt.strip().startswith(("#", "|", ">")) and nxt.startswith((" ", "\t")):
                text += " " + nxt.strip()
                i += 1
            else:
                break
        ordered = marker[0].isdigit()
        style = "List Number" if ordered else "List Bullet"
        if len(indent) >= 2:
            style += " 2"
        p = doc.add_paragraph(style=style)
        add_runs(p, text)
        i += 1
        continue

    # --- paragraph (join wrapped lines) ---
    text = stripped
    while i + 1 < len(lines):
        nxt = lines[i + 1].strip()
        if nxt and not nxt.startswith(("#", "|", ">", "-", "*")) and not re.match(r"^\d+\.\s", nxt) and nxt != "---":
            text += " " + nxt
            i += 1
        else:
            break
    add_runs(doc.add_paragraph(), text)
    i += 1

doc.save(OUT)
print("Saved:", OUT)
if missing_figures:
    print("MISSING (placeholders inserted) - drop these into docs/images/ and re-run:")
    for f in missing_figures:
        print("   -", f)
else:
    print("All figures embedded.")
